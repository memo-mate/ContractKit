import mammoth
from bs4 import BeautifulSoup
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import parse_xml
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from llama_index.core.bridge.pydantic import BaseModel, ConfigDict, Field


def runs_merge(paragraph: Paragraph) -> Run | None:
    """
    Merge all runs in a paragraph into a single run.

    Args:
        paragraph (_Paragraph): The paragraph to merge runs in.

    Returns:
        Optional[_Run]: The merged run, or None if there are no runs.
    """
    runs = paragraph.runs
    if len(runs) == 0:
        runs = list(Run(r, paragraph) for r in parse_xml(paragraph._element.xml.replace("fld", "r")).r_lst)
    if len(runs) == 1:
        return runs[0]
    if len(runs) == 0:
        return None

    run = max(runs, key=lambda x: len(x.text))
    run.text = paragraph.text
    for r in runs:
        if r != run:
            r._r.getparent().remove(r._r)
    return run


def count_placeholders(text: str) -> int:
    """
    Count the number of placeholders in a text.
    """
    count = 0
    start = 0
    while True:
        start = text.find("<|", start)
        if start == -1:
            break
        end = text.find("|>", start)
        if end == -1:
            break
        count += 1
        start = end + 2
    return count


def get_html_tables(word_path: str) -> list[str]:
    """
    Get the html tables from a word file.

    Args:
        word_path (str): The path to the word file.

    Returns:
        list[str]: The html tables.
    """
    assert word_path.endswith(".docx"), "word_path must be a .docx file"

    with open(word_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
    html_content = result.value

    soup = BeautifulSoup(html_content, "html.parser")

    tables = soup.find_all("table")

    return [table.prettify() for table in tables] # type: ignore


class Content(BaseModel):
    """Segment of the contract"""
    model_config = ConfigDict(arbitrary_types_allowed=True)

    id: int
    content_type: str
    content: str
    paragraphs: list[Paragraph] = Field(default_factory=list, exclude=True)
    raw: Paragraph | Table = Field(exclude=True)


# def parse_table(table: Table) -> tuple[str, list[Paragraph]]:
#     """
#     Parse a table into a markdown table and a list of paragraphs.
#     """
#     return "", []


def get_contents(document_path: str) -> tuple[list[Content], DocxDocument]:
    """
    Get the contents of a document.
    """
    document = Document(document_path)
    tables = get_html_tables(document_path)

    content_id = 0
    table_id = 0
    contents = []
    for element in document.elements:
        if isinstance(element, Paragraph):
            contents.append(
                Content(
                    id=content_id,
                    content_type="paragraph",
                    content=element.text,
                    paragraphs=[element],
                    raw=element,
                )
            )
        elif isinstance(element, Table):
            paragraphs = []
            for row in element.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)
            contents.append(
                Content(
                    id=content_id,
                    content_type="table",
                    content=tables[table_id],
                    paragraphs=paragraphs,
                    raw=element,
                )
            )
            table_id += 1
        content_id += 1
    return contents, document


def set_paragraph_text(paragraph: Paragraph, text: str) -> Run:
    """
    Set the text of a paragraph.

    Args:
        paragraph (_Paragraph): The paragraph to set the text of.
        text (str): The text to set.

    Returns:
        _Run: The run that was set.
    """
    run = runs_merge(paragraph)
    if run is None:
        run = paragraph.add_run(text)
    else:
        run.text = text
    return run


def cell_paragraphs_merge(cell: _Cell) -> Paragraph:
    """
    Merge all paragraphs in a cell into a single paragraph.

    Args:
        cell (_Cell): The cell to merge paragraphs in.

    Returns:
        Paragraph: The merged paragraph.
    """
    paragraphs = cell.paragraphs
    if len(paragraphs) == 0:
        paragraph = cell.add_paragraph()
        return paragraph
    if len(paragraphs) == 1:
        return paragraphs[0]
    paragraph = max(paragraphs, key=lambda x: len(x.text))
    set_paragraph_text(paragraph, cell.text)
    for p in paragraphs:
        if p != paragraph:
            p._p.getparent().remove(p._p)
    return paragraph

def remove_row(row: _Row) -> None:
    row._element.getparent().remove(row._element)

def remove_cell(cell: _Cell) -> None:
    cell._element.getparent().remove(cell._element)
