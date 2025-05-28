from typing import Any, List, Optional

from docx.document import Document
# from docx.text.paragraph import Paragraph
# from docx.text.run import Run
from llama_index.core.bridge.pydantic import Field
from llama_index.core.llms import LLM, ChatMessage
from llama_index.core.memory import ChatMemoryBuffer
from llama_index.core.output_parsers.utils import parse_json_markdown
from llama_index.core.prompts import PromptTemplate
from llama_index.core.settings import Settings
from llama_index.core.tools.types import BaseTool
from llama_index.core.workflow import (
    Context,
    Event,
    StartEvent,
    StopEvent,
    Workflow,
    step,
)

from prompts.prompts import filling_prompt, handle_incomplete_prompt
from .utils import Content, get_contents, set_paragraph_text


class InputEvent(StartEvent):
    # document: Document = Field(description="The document to fill")
    document_path: str = Field(description="The path of the document to fill")


class ParagraphEvent(Event):
    contents: List[Content] = Field(default_factory=list, description="The contents to fill")
    document: Document = Field(description="The document to fill", exclude=True)

class FillingEvent(Event):
    name: str
    msg: str
    data: Any = Field(default=None, exclude=True)

    @property
    def message(self) -> str:
        return self.msg

    @message.setter
    def message(self, value: str) -> None:
        self.msg = value


class IncompleteEvent(Event):
    written: str = Field(
        default="",
        description="The contents that have been filled",
    )
    incomplete_contents: List[str] = Field(
        default_factory=list,
        description="The contents that still has placeholder",
    )
    document: Document = Field(description="The document to fill", exclude=True)

class FillingAgent(Workflow):
    def __init__(
        self,
        *args: Any,
        llm: LLM | None = None,
        chat_history: Optional[List[ChatMessage]] = None,
        tools: List[BaseTool] | None = None,
        system_prompt: PromptTemplate | None = None,
        handle_incomplete_prompt: PromptTemplate = handle_incomplete_prompt,
        verbose: bool = False,
        timeout: float = 360.0,
        name: str = "Filling",
        write_events: bool = True,
        description: str | None = None,
        **kwargs: Any,
    ) -> None:
        kwargs.update({"verbose": verbose, "timeout": timeout})
        super().__init__(*args, **kwargs)
        self.tools = tools or []
        self.name = name
        self.write_events = write_events
        self.description = description

        self.llm = llm or Settings.llm

        self.system_prompt = system_prompt or filling_prompt
        self.handle_incomplete_prompt = handle_incomplete_prompt

        self.memory = ChatMemoryBuffer.from_defaults(llm=self.llm, chat_history=chat_history)

    @step
    async def prepare_paragraphs(self, ctx: Context, evt: InputEvent) -> ParagraphEvent:

        await ctx.set("written", [])
        await ctx.set("incomplete_contents", [])
        document_path = evt.document_path
        contents, doc = get_contents(document_path)
        if self.write_events:
            ctx.write_event_to_stream(FillingEvent(name=self.name, msg="Preparing paragraphs", data=contents))
        return ParagraphEvent(contents=contents, document=doc)

    @step
    async def fill(self, ctx: Context, evt: ParagraphEvent) -> IncompleteEvent:
        for content in evt.contents:
            messages = self.memory.get().copy()
            text = content.content
            # if include placeholder, then fill
            if "<|" in text and "|>" in text:
                messages.append(ChatMessage(role="assistant", content="接下来请传入带有占位符的当前段落。"))
                messages.append(ChatMessage(role="user", content=text))
                written = await ctx.get("written")
                context = "\n".join(written)
                system_msg = ChatMessage(role="system", content=self.system_prompt.format(context=context))
                messages.insert(0, system_msg)

                response = await self.llm.achat(
                    messages=messages,
                )
                response_text = response.message.content
                if response_text is not None:
                    json_data = parse_json_markdown(response_text)
                    # assert len(json_data) >= len(content.paragraphs), (
                    #     f"json_data: {json_data}, content.paragraphs: {[p.text for p in content.paragraphs]}"
                    # )
                    if json_data is None:
                        continue
                    for paragraph in content.paragraphs:
                        for key, value in json_data.items():
                            if value is not None:
                                # if placeholder is in the paragraph, then replace it
                                if key in paragraph.text:
                                    new_paragraph_text = paragraph.text.replace(key, value)
                                    text = text.replace(key, value)
                                    set_paragraph_text(paragraph, new_paragraph_text)
                # if still has placeholder, then add to incomplete_content
                if "<|" in text and "|>" in text:
                    ctx.data["incomplete_contents"].append(text)
            ctx.data["written"].append(text)
            if self.write_events:
                ctx.write_event_to_stream(FillingEvent(name=self.name, msg="Filling", data=text))
        written = await ctx.get("written")
        incomplete_contents = await ctx.get("incomplete_contents")
        if self.write_events:
            ctx.write_event_to_stream(FillingEvent(name=self.name, msg="Filling", data=incomplete_contents))
        return IncompleteEvent(
            written="\n".join(written),
            incomplete_contents=incomplete_contents,
            document=evt.document
        )

    @step
    async def handle_incomplete_contents(self, ctx: Context, evt: IncompleteEvent) -> StopEvent:
        written_text = evt.written
        incomplete_contents = evt.incomplete_contents
        if len(incomplete_contents) == 0:
            return StopEvent(result={"response": "已完成合同撰写。", "document": evt.document})
        response = await self.llm.apredict(self.handle_incomplete_prompt, written=written_text)
        if self.write_events:
            ctx.write_event_to_stream(FillingEvent(name=self.name, msg="Handling incomplete contents", data=response))
        return StopEvent(result={"response": response, "document": evt.document})
